"""生成 word"""
# import os
import decimal
# from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from web.enum.enum import Tag,TagString, get_account_name,get_consumer_name


def to_decimal_if_is_string(data):
    """

    :return:
    """
    return decimal.Decimal(data) if isinstance(data, str) else data


def set_style(document):
    """
    设置样式
    :return:
    """
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(14)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)


def word_head2(document, title):
    # 二级标题
    head2 = document.add_heading(level=2)
    head2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = head2.add_run(title)
    # run.font.name=u'宋体'
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 段落后行距
    head2.paragraph_format.space_after = Pt(30)

def word_head3(document, title):
    # 三级标题
    head3 = document.add_heading(level=3)
    run = head3.add_run(title)
    # run.font.name=u'宋体'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 0)


def creat_demo(document):
    # 报告标题
    head = document.add_heading()
    head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    head.paragraph_format.space_before = Pt(36)
    run = head.add_run(u"生成的报告")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 标题颜色
    document.styles['Normal'].font.name = u'微软雅黑'  # 设置西文字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体使用字体2->宋体
    # 添加分页
    document.add_page_break()

    # 二级标题
    head2 = document.add_heading(level=2)
    head2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = head2.add_run(u'账单分析')
    # run.font.name=u'宋体'
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 段落后行距
    head2.paragraph_format.space_after = Pt(30)

    # 二级级标题段落
    # document.styles['Normal'].font.name = u'宋体'
    p = document.add_paragraph()
    run = p.add_run(u'''    段落1
        段落2''')
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(15)
    # 单倍行距
    p.paragraph_format.line_spacing = Pt(30)
    # 段落后行距
    p.paragraph_format.space_after = Pt(30)

    # 三级标题
    head3 = document.add_heading(level=3)
    run = head3.add_run(u'分析对象:')
    # run.font.name=u'宋体'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 段落后行距
    head2.paragraph_format.space_after = Pt(30)

    # 设置中文字体
    # run = paragraph.add_run(u'设置中文字体，')
    # run.font.name=u'宋体'
    # r = run._element
    # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'MD5:  ')
    run.bold = True
    run.font.size = Pt(12)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'TYPE:  ')
    run.bold = True
    run.font.size = Pt(12)
    paragraph = document.add_paragraph()
    # 设置粗体小四
    run = paragraph.add_run(u'文件名称:  ')
    run.bold = True
    run.font.size = Pt(12)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'文件大小:  ')
    run.bold = True  # 加粗
    run.font.size = Pt(12)  # 小四

    # 报告1
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'判定:  ')
    run.bold = True
    run.font.size = Pt(12)

    # 报告2
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'报告2:  ')
    run.bold = True
    run.font.size = Pt(12)

    # 分析总结
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'分析结果:  ')
    run.bold = True
    run.font.size = Pt(12)

    # 检测结果
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'检测结果:  ')
    run.bold = True
    run.font.size = Pt(12)

    # 关系
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'关系:  ')
    run.bold = True  # 加粗
    run.font.size = Pt(12)  # 小四

    # 关键字符串
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'关键字符串:  ')
    run.bold = True
    run.font.size = Pt(12)

    head4 = document.add_heading(level=3)
    run = head4.add_run(u'附件:')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # 注1
    head5 = document.add_heading(level=3)
    run = head5.add_run(u'注1:')
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(65, 186, 126)
    head5.paragraph_format.line_spacing = Pt(30)

    p = document.add_paragraph()
    p.paragraph_format.line_spacing = Pt(30)
    run = p.add_run(u'''    注解1段落''')
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(12)

    # 注2
    head6 = document.add_heading(level=3)
    run = head6.add_run(u'注2:')
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(65, 186, 126)
    head6.paragraph_format.line_spacing = Pt(30)

    p = document.add_paragraph()
    p.paragraph_format.line_spacing = Pt(30)
    run = p.add_run(u'''    注解2段落''')
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(12)


def first_page(document):
    # head
    head = document.add_heading()
    head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    head.paragraph_format.space_before = Pt(36)
    run = head.add_run(u"三月月报")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 标题颜色
    document.styles['Normal'].font.name = u'微软雅黑'  # 设置西文字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体使用字体2->宋体
    document.add_page_break()


def create_doc(data):
    """
    :param data:
    :return:
    """
    document = Document()

    # set_style(document)


    consumer_info = get_consumer(data)

    summary_data = paragraph_summary(data)

    tag_txt = get_tag_txt(data)

    first_page(document)

    second_page(document, data, summary_data)

    third_page(document, summary_data, consumer_info, tag_txt)

    # creat_demo(document)
    # cwd = os.getcwd()
    # now = datetime.now()
    # dt_string = now.strftime("%d/%m/%Y %H:%M")
    # file_name = "\\demo\\{}.docx".format(dt_string)
    document.save('demo.docx')


def get_tag_name(data):
    """

    :param data:
    :return:
    """
    return TagString[Tag(data).name]


def second_page(document, data, render_data):
    """

    :param document:
    :param render_data:
    :param data:
    :return:
    """

    word_head2(document, u'概览')
    # document.add_heading('标题1', level=1)

    render_summary(document, render_data)

    paragraph_account_type(document, data)

    paragraph_consumer(document, data)

    paragraph_tag(document, data)

    document.add_page_break()


def render_summary(document, render_data):
    """

    :param document:
    :param render_data:
    :return:
    """
    document.add_paragraph(render_data.get("total_txt"))


    word_head3(document, u'分类')

    document.add_paragraph(render_data.get("detail_current_month"))


def paragraph_tag(document, data):
    """
    消费来源：
    固定消费： 3000元， 变动消费：3300，
    :param document:
    :param data:
    :return:
    """
    detail = ""
    for item in data['tag']:
        label = get_tag_name(item['tag'])
        cost = item['total_amount']
        percent = item['percent']
        detail = detail + f"【{label}】{cost}元，占比 {percent}%。"
    word_head3(document, u'标签')
    document.add_paragraph(detail)
    document.add_paragraph('')


def get_tag_txt(data):
    """

    :param document:
    :param data:
    :return:
    """
    detail = ""
    compare_last_month_more = ""
    compare_last_month_less = ""
    for item in data['tag']:
        label = get_tag_name(item['tag'])
        tag_id = item['tag']
        cost_last_month = data.get('tag_last_month_dict').get(tag_id).get('total_amount')
        cost = item['total_amount']
        percent = item['percent']
        compare_last_month = compare_sentence(cost, cost_last_month)
        if cost > cost_last_month:
            compare_last_month_more = compare_last_month_more + f"【{label}】对比上月支出{compare_last_month}, "
        else:
            compare_last_month_less = compare_last_month_less + f"【{label}】对比上月支出{compare_last_month}, "
        detail = detail + f"【{label}】{cost}元，占比 {percent}%。"

    return {
        "compare_last_month_more": compare_last_month_more,
        "compare_last_month_less": compare_last_month_less
    }


def paragraph_account_type(document, data):
    """
    钱包支出
    :param document:
    :param data:
    :return:
    """
    detail = ""
    for item in data['account']:
        label = get_account_name(item['account_type'])
        cost = item['total_amount']
        detail = detail + f"【{label}】{cost}元，"
    # document.add_paragraph("钱包支出：")
    # 三级标题
    head3 = document.add_heading(level=3)
    run = head3.add_run(u'钱包支出')
    # run.font.name=u'宋体'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 0)
    document.add_paragraph(detail)


def get_consumer(data):
    """

    :return:
    """
    detail = ""
    compare_last_month_more = ""
    compare_last_month_less = ""
    for item in data['consumer']:
        label = get_consumer_name(item['consumer'])
        cost = item['total_amount']
        percent = item['percent']
        last_total = data.get('consumer_last_month_dict').get(item['consumer']).get('total_amount')
        compare = compare_sentence(cost, last_total)
        if cost > last_total:
            compare_last_month_more = compare_last_month_more + f"【{label}】对比上月支出{compare}, "
        else:
            compare_last_month_less = compare_last_month_less + f"【{label}】对比上月支出{compare}, "
        detail = detail + f"【{label}】{cost}元，占比 {percent}%。"

    return {
        "consumer_detail": detail,
        "compare_last_month_more": compare_last_month_more,
        "compare_last_month_less": compare_last_month_less,
    }


def paragraph_consumer(document, data):
    """
    成员支出
    :param document:
    :param data:
    :return:
    """
    detail = ""
    for item in data['consumer']:
        label = get_consumer_name(item['consumer'])
        cost = item['total_amount']
        percent = item['percent']
        detail = detail + f"【{label}】{cost}元，占比 {percent}%。"
    word_head3(document, u'成员支出')
    document.add_paragraph(detail)


def paragraph_summary(data):
    """
    本月总支出 【3333】 元。
    其中包括【食品吃喝】【3331】元，购物消费 【2】元。
    :param document:
    :param data:
    :return:
    """
    detail = "其中包括"
    compare_detail_cost_txt_more = ""
    compare_detail_cost_txt_less = ""
    total = decimal.Decimal(0)

    category_last_month_obj = data['category_last_month_obj']
    for item in data['category']:
        title = item['label']
        item_id = item['id']
        cost = item['amount']
        total += decimal.Decimal(item['amount'])
        detail = detail + f"【{title}】{cost}元，"
        if category_last_month_obj.get(item_id):
            last_item = category_last_month_obj.get(item_id)

            amount_now = decimal.Decimal(item['amount']) if isinstance(item['amount'], str) else item['amount']
            amount_last = decimal.Decimal(last_item['amount']) if isinstance(last_item['amount'], str) else last_item['amount']

            compare_detail_cost = compare_sentence(amount_now, amount_last)
            if amount_now > amount_last:
                compare_detail_cost_txt_more = compare_detail_cost_txt_more + f"【{title}】对比上月支出{compare_detail_cost}, "
            else:
                compare_detail_cost_txt_less = compare_detail_cost_txt_less + f"【{title}】对比上月支出{compare_detail_cost}, "

    total_txt = f"本月总支出 {total} 元。"

    detail_current_month = detail

    total_last_month = decimal.Decimal(0)

    for item in data['category_last_month']:
        total_last_month += decimal.Decimal(item.get('amount'))

    last_year_month_cost_avg_txt = compare_sentence(total, data.get("last_year_month_cost_avg")[0].get("month_avg"))
    avg_of_last_quarter_amount_txt = compare_sentence(total, data.get("avg_of_last_quarter_amount")[0].get("month_avg"))
    compare_last_month_total = compare_sentence(total, total_last_month)

    compare_last_total_txt = f"总支出对比上月 {compare_last_month_total} ，"
    compare_last_year_month_cost_avg_txt = f" 对比去年每月平均{last_year_month_cost_avg_txt}"
    compare_avg_of_last_quarter_amount_txt = f" 对比上个季度每月平均{avg_of_last_quarter_amount_txt}"

    return {
        'total_txt': total_txt,
        'compare_last_total_txt': compare_last_total_txt,
        'detail_current_month': detail_current_month,
        'compare_detail_cost_txt_more': compare_detail_cost_txt_more,
        'compare_detail_cost_txt_less': compare_detail_cost_txt_less,
        'last_year_month_cost_avg_txt': compare_last_year_month_cost_avg_txt,
        'compare_avg_of_last_quarter_amount_txt': compare_avg_of_last_quarter_amount_txt,
    }


def third_page(document, render_data, consumer_info, tag_txt):
    """

    :param document:
    :param render_data:
    :param consumer_info:
    :param tag_txt:
    :return:
    """
    word_head2(document, u'各项支出对比')

    render_summary_compare(document, render_data)
    render_consumer_compare(document, consumer_info)
    render_tag_compare(document, tag_txt)


def render_consumer_compare(document, render_data):
    word_head3(document, u'成员支出')
    document.add_paragraph("支出增加:")
    document.add_paragraph(render_data.get('compare_last_month_more'))
    document.add_paragraph("支出减少:")
    document.add_paragraph(render_data.get('compare_last_month_less'))


def render_tag_compare(document, render_data):
    word_head3(document, u'标签')
    document.add_paragraph("支出增加:")
    document.add_paragraph(render_data.get('compare_last_month_more'))
    document.add_paragraph("支出减少:")
    document.add_paragraph(render_data.get('compare_last_month_less'))


def render_summary_compare(document, render_data):
    """

    :param document:
    :param render_data:
    :return:
    """
    para = (render_data.get("compare_last_total_txt")
            + render_data.get("last_year_month_cost_avg_txt")
            + render_data.get("compare_avg_of_last_quarter_amount_txt"))

    document.add_paragraph(para)
    word_head3(document, u"分类")
    document.add_paragraph("支出增加:")
    document.add_paragraph(render_data.get('compare_detail_cost_txt_more'))
    document.add_paragraph("支出减少:")
    document.add_paragraph(render_data.get('compare_detail_cost_txt_less'))


def compare_sentence(cur, old):
    """
    对比
    :param cur:
    :param old:
    :return:
    """
    sentence = ''
    a = to_decimal_if_is_string(cur)
    b = to_decimal_if_is_string(old)

    cost = round(a - b, 2)
    if cost > 0:
        sentence = f"增加了{cost}元"
    elif cost == 0:
        sentence = "不变"
    else:
        sentence = f"减少了{abs(cost)} 元 "

    return sentence
